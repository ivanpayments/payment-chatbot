"""Rebuild page_text.docx from scratch with three pages:

    1. Chatbot landing-page copy    (text preserved byte-identical)
    2. Portfolio landing page copy  (text preserved + Case 02 appended)
    3. Routing Simulator copy       (new, sourced from Project 2 content.json)

Styling rules applied for visual consistency:
    - Title style: Word "Title" on the first line of each page
    - Section labels ([HERO], [SECTION - ...], etc.): Word "Heading 2"
    - Field labels (Uptitle, H1, Body, etc.): Normal, italic, gray run
    - Value text: Normal
    - Hard page break between pages
    - One-line editor instruction at top of each page

CRITICAL: the chatbot section labels and field-label prefixes are parsed by
Claude files/sync_docx.py. They must stay byte-identical. The script below
uses the exact strings found in the live docx.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "page_text.docx"

EDITOR_INSTRUCTION = (
    'Edit the text below. Tell Claude "the doc is updated" and the page will '
    "be synced. Do not change section labels in square brackets \u2014 they are keys."
)


def add_page_title(doc, text):
    p = doc.add_paragraph(text, style="Title")
    return p


def add_instruction(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.size = Pt(10)
    return p


def add_section_label(doc, text):
    """Bracketed section label, style = Heading 2 so sync_docx.py can find it."""
    p = doc.add_paragraph(text, style="Heading 2")
    return p


def add_field_label(doc, text):
    """Instruction-ish label paragraph (e.g. 'Uptitle (top small line, all caps)')."""
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p


def add_value(doc, text, style="Normal"):
    p = doc.add_paragraph(text, style=style)
    return p


def add_blank(doc):
    doc.add_paragraph()


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


# ---------------------------------------------------------------------------
# Chatbot page content (verbatim from existing docx)
# ---------------------------------------------------------------------------

def build_chatbot_page(doc):
    # NOTE: section_heading for chatbot page uses Heading 1 because sync_docx.py
    # matches `p.style.name == "Heading 1"`. We must keep Heading 1 for these.
    def sec(text):
        doc.add_paragraph(text, style="Heading 1")

    add_page_title(doc, "Chatbot landing-page copy")
    add_instruction(doc, EDITOR_INSTRUCTION)

    # [HERO]
    sec("[HERO]")
    add_field_label(doc, "Uptitle (top small line, all caps)")
    add_value(doc, "Payment analytics")
    add_field_label(doc, "H1 (main title)")
    add_value(doc, "A conversational analyst for the Head of Payments")
    add_field_label(doc, "Stats (one per line, format: LABEL | VALUE)")
    add_value(doc, "Countries  | 30")
    add_value(doc, "Annual TPV | ~$2B")
    add_value(doc, "Acquirers | 14")
    add_field_label(doc, "Subtitle (below title)")
    add_value(
        doc,
        "Imagine you are the Head of Payments at a global SaaS company. "
        "The chatbot helps you manage the business by answering every "
        "question about your transactions in plain English",
    )

    # [SECTION - WHAT THIS IS]
    sec("[SECTION \u2014 WHAT THIS IS]")
    add_field_label(doc, "Label")
    add_value(doc, "What is this")
    add_field_label(doc, "Heading")
    add_value(
        doc,
        "A chatbot leveraging transaction data to answer questions about your "
        "payment performance",
    )
    add_field_label(doc, "Body")
    add_value(
        doc,
        "Chatbot leverages transaction data to conduct performance analysis of "
        "payments setup and answer ad-hoc questions from various stakeholders "
        "(CFO, head or risk, operations lead etc.). Chatbot is accessible on "
        "web, via Whatsapp or SMS.",
    )

    # [SECTION - WHY THIS MATTERS]
    sec("[SECTION \u2014 WHY THIS MATTERS]")
    add_field_label(doc, "Label")
    add_value(doc, "Why this matters")
    add_field_label(doc, "Heading")
    add_value(doc, "Analysis of payments performance takes days of manual work")
    add_field_label(doc, "Body")
    add_value(
        doc,
        "Merchants of similar scale typically dedicate three to five analysts "
        "to monitor their payments setup. Time-consuming activities include "
        "among others isolating reasons for approval drop across providers, "
        "tracking 200+ routing rules in Excel, identifying PSP fee changes "
        "which were not communicated, and detecting provider outages and "
        "degradation. This chatbot turns a multi-day investigation into a "
        "two-minute conversation",
    )

    # [SECTION - WHO IT IS FOR]
    sec("[SECTION \u2014 WHO IT IS FOR]")
    add_field_label(doc, "Label")
    add_value(doc, "Who are the users")
    add_field_label(doc, "Heading")
    add_value(doc, "The Head of Payments and their team")
    add_field_label(doc, "Body")
    add_value(
        doc,
        "The primary users are the Head of Payments, their team, and "
        "stakeholders in the organization interested in payment performance: "
        "the CFO interested in cost of payments; the Chief Commercial Officer "
        "who would like to understand revenue lost due to gaps in payment "
        "methods offering; the fraud lead monitoring fraud rates; the regional "
        "operations team resolving customer payment complaints; the strategy "
        "team evaluating acquirer rationalisation",
    )

    # [SECTION - WHY IT EXISTS]
    sec("[SECTION \u2014 WHY IT EXISTS]")
    add_field_label(doc, "Label")
    add_value(doc, "What it answers")
    add_field_label(doc, "Heading")
    add_value(
        doc,
        "Every question about payments, from long-term contract economics to "
        "real-time performance across providers, countries and payment methods",
    )
    add_field_label(doc, "Body (intro paragraph before problems)")
    add_value(
        doc,
        "The Head of Payments can ask the chatbot any question based on the "
        "transaction data. The chatbot can also generate files such as "
        "financial models, voice notes, PDF reports. There are several "
        "recurrent problems with payments that are typical for companies of "
        "this scale:",
    )
    add_field_label(
        doc,
        "Points (one per line, bold the opening phrase by ending with a period "
        "+ space before the rest)",
    )
    add_value(
        doc,
        "Volume commitments and rebate tiers are tracked only on a quarterly "
        "basis. Contracts carry minimums, step-downs and rebate thresholds "
        "which can be leveraged if data is available and routing decisions are "
        "made on a weekly basis. Providers increase fees without notice, or "
        "make mistakes in billing. The chatbot can monitor the commitments and "
        "suggest the best actions to optimize the payments setup",
    )
    add_value(
        doc,
        "Approval-rate performance across the stack is opaque. Approval rates "
        "can diverge by 5\u201315 pp across the acquirer portfolio. Merchant "
        "teams often notice outages and provider unavailability only after "
        "customers complain about failed transactions. The chatbot can alert "
        "users of ongoing outages",
    )
    add_value(
        doc,
        "Local payment methods and checkout configuration are often not "
        "optimized. Each country\u2019s customers expect specific payment "
        "methods, currencies and installment options. The chatbot can analyze "
        "customer purchase patterns and checkout configuration, and identify "
        "opportunities for improvement",
    )

    # [SECTION - TRANSACTIONS]
    sec("[SECTION \u2014 TRANSACTIONS]")
    add_field_label(doc, "Label")
    add_value(doc, "Which data it uses")
    add_field_label(doc, "Heading")
    add_value(
        doc,
        "Synthetic transaction data, designed to reproduce the statistical "
        "behaviour of payments within a global SaaS company",
    )
    add_field_label(doc, "Body paragraph 1 (after stats)")
    add_value(
        doc,
        "The data includes various types of transaction data typically "
        "handled by SaaS companies:",
    )
    add_value(doc, "Authorization and processing status", style="List Paragraph")
    add_value(doc, "Recurrent subscriptions", style="List Paragraph")
    add_value(doc, "3DS flow and tokenization", style="List Paragraph")
    add_value(
        doc,
        "Per-transaction FX, interchange and scheme fees applied",
        style="List Paragraph",
    )
    add_value(doc, "Retries and provider routing", style="List Paragraph")
    add_field_label(doc, "Body paragraph 2 (synthetic data note)")
    add_value(
        doc,
        "The dataset is synthetic. It has been engineered to reproduce the "
        "statistical behaviour of a book at this scale and is used for "
        "pattern analysis. In a production deployment, the same chatbot can "
        "be rebuilt on top of the company\u2019s real systems and data "
        "sources (transaction processing engine, acquirer statements, data "
        "warehouse, etc.)",
    )
    add_field_label(
        doc,
        'Note: to make the first two words bold ("The dataset is synthetic."), '
        "leave them as the opening of the paragraph. Claude handles the "
        "formatting when syncing.",
    )

    # [SECTION - HOW IT IS BUILT]
    sec("[SECTION \u2014 HOW IT IS BUILT]")
    add_field_label(doc, "Label")
    add_value(doc, "How it works")
    add_field_label(doc, "Heading")
    add_value(doc, "Claude Opus writing scripts in a sandboxed environment")
    add_field_label(doc, "Body")
    add_value(
        doc,
        "Claude Opus 4.6 writes and runs Python scripts in a sandboxed "
        "environment against the transaction dataset, and streams the answer "
        "back. Two selected questions carry pre-generated answers to enable "
        "quick user demo. The same agent also answers on WhatsApp and SMS, "
        "connected via Twilio",
    )

    # [CALL TO ACTION]
    sec("[CALL TO ACTION]")
    add_field_label(doc, "Button text")
    add_value(doc, "Give it a try")

    # [CHAT PANEL]
    sec("[CHAT PANEL]")
    add_field_label(doc, "Title")
    add_value(doc, "Chat with co-pilot")
    add_field_label(doc, "Subtitle")
    add_value(doc, "Ask anything about your transactions")
    add_field_label(doc, "Chips label (above suggested questions)")
    add_value(doc, "Try a question from your team")

    # [SUGGESTED QUESTIONS]
    sec("[SUGGESTED QUESTIONS]")
    add_field_label(doc, "One per line, format: STAKEHOLDER | question text")
    add_value(
        doc,
        "CFO | Cost per approved transaction jumped from $43 to $63 last "
        "quarter. What's driving it, and what's the single biggest lever we "
        "can pull this quarter to solve this?",
    )
    add_value(
        doc,
        "LATAM ops | Brazil approvals fell 3 percentage points this week. "
        "What caused it and how fast can we fix it?",
    )
    add_value(
        doc,
        "Product | How much revenue has the card-refresh tool saved us so "
        "far, and is it worth rolling it out everywhere?",
    )
    add_value(
        doc,
        "Strategy | When a payment fails at one processor, trying a different "
        "one often works. Which five markets would recover the most money "
        "from that?",
    )
    add_value(
        doc,
        "Finance | Which of our processors is charging more than the "
        "contracted rate for currency conversion, and how much is it costing "
        "us a year?",
    )
    add_value(
        doc,
        "CEO brief | How does our approval rate compare to other SaaS "
        "companies of our size, and what is the single biggest change that "
        "would close the gap?",
    )

    # [CHAT UI LABELS]
    sec("[CHAT UI LABELS]")
    add_field_label(doc, "Input placeholder")
    add_value(doc, "Type a question\u2026")
    add_field_label(doc, "Send button")
    add_value(doc, "Send")
    add_field_label(doc, '"You" label on user messages')
    add_value(doc, "You")
    add_field_label(doc, '"Claude" label on bot messages')
    add_value(doc, "Claude")
    add_field_label(doc, "Thinking label (while waiting for response)")
    add_value(
        doc,
        "Claude is thinking\u2026 Response time is typically 2-3 minutes as "
        "Claude needs to analyze source files to produce the most "
        "comprehensive answer. Try one of the questions with pre-generated "
        "answers for faster response",
    )


# ---------------------------------------------------------------------------
# Portfolio page (verbatim + new Case 02 block)
# ---------------------------------------------------------------------------

def build_portfolio_page(doc):
    add_page_title(doc, "Portfolio landing page copy")
    add_instruction(doc, EDITOR_INSTRUCTION)

    # [PORTFOLIO - HERO]
    add_section_label(doc, "[PORTFOLIO \u2014 HERO]")
    add_field_label(doc, "Uptitle (top small line, all caps)")
    add_value(doc, "Portfolio ")
    add_field_label(doc, "H1 (main title)")
    add_value(doc, "Payments with Ivan")
    add_field_label(doc, "Subtitle (below title)")
    add_value(
        doc,
        "Ivan Antonov \u2014 payments from strategy to working software\n"
        "Background\n"
        "- Digital transformation at ING\n"
        "- Advisory to financial institutions at McKinsey\n"
        "- Agent engineering at Yuno",
    )

    # [PORTFOLIO - CASE 01]
    add_section_label(doc, "[PORTFOLIO \u2014 CASE 01]")
    add_field_label(doc, "Tab label")
    add_value(doc, "Case 01 \u2014 Payment Analytics Chatbot")
    add_field_label(doc, "Heading")
    add_value(doc, "A payment analytics chatbot")
    add_field_label(doc, "Lede (first paragraph, slightly larger)")
    add_value(
        doc,
        "What is it? A natural-language chatbot answering questions about "
        "transaction data. Merchant payment teams can ask questions and get "
        "answers in seconds, without the need to rely on other teams",
    )
    add_value(
        doc,
        "What problem it solves? Payment teams get real-time visibility on "
        "provider outages, costs and approval rates, as well as "
        "recommendations on actions to take \u2013 something that required a "
        "few days of analysis before\n\n"
        "How it works? A synthetic dataset with 100k transactions, "
        "reproducing the statistical patterns of a book at this scale. The "
        "chatbot runs an LLM query across the dataset. In a production "
        "setting, the chatbot can be connected to the company\u2019s systems "
        "and data sources",
    )
    add_field_label(doc, "Link text")
    add_value(doc, "Read more and try it \u2192")
    add_field_label(doc, "Link URL")
    add_value(doc, "https://www.ivanantonov.com/chatbot/")

    # [PORTFOLIO - CASE 02]  (NEW)
    add_section_label(doc, "[PORTFOLIO \u2014 CASE 02]")
    add_field_label(doc, "Tab label")
    add_value(doc, "Case 02 \u2014 Payment Routing Simulator")
    add_field_label(doc, "Heading")
    add_value(doc, "A payment routing simulator")
    add_field_label(doc, "Lede 1 (What is it?)")
    add_value(
        doc,
        "What is it? A multi-archetype mock payment gateway modeling five "
        "processor archetypes \u2013 global-acquirer, regional-bank-processor, "
        "APM-specialist, cross-border-FX-specialist, and high-risk "
        "orchestrator \u2013 with realistic ISO 8583 decline codes, 3DS v2.2 "
        "flows, and HMAC-signed webhooks",
    )
    add_field_label(doc, "Lede 2 (What problem does it solve?)")
    add_value(
        doc,
        "What problem does it solve? Sandbox environments approve on magic "
        "cards; production behaves differently. Merchant integration teams "
        "cannot test routing logic against realistic declines until customers "
        "complain. The simulator returns what merchants actually see in "
        "production, so integrations ship correct the first time",
    )
    add_field_label(doc, "Lede 3 (How does it work?)")
    add_value(
        doc,
        "How does it work? Python with FastAPI, a Postgres payment state "
        "machine, Redis idempotency cache, a Kafka event bus, and Celery "
        "webhook workers. Installable via pip install payment-routing-simulator, "
        "live API under /routing-simulator/api/, and wired into Case 01\u2019s "
        "chatbot as a routing-recommendation tool",
    )
    add_field_label(doc, "Link text")
    add_value(doc, "Read more \u2192")
    add_field_label(doc, "Link URL")
    add_value(doc, "/routing-simulator/")

    # [PORTFOLIO - QUOTE]
    add_section_label(doc, "[PORTFOLIO \u2014 QUOTE]")
    add_field_label(doc, "Large quote on gray block")
    add_value(
        doc,
        "I help financial institutions, fintechs and merchant payment teams "
        "build and optimize payments, from strategy to working software",
    )

    # [PORTFOLIO - FOOTER]
    add_section_label(doc, "[PORTFOLIO \u2014 FOOTER]")
    add_field_label(doc, "Tagline")
    add_value(doc, "Ivan Antonov")
    add_field_label(doc, "Email")
    add_value(doc, "ivan.antonov@outlook.com")
    add_field_label(doc, "Website")
    add_value(doc, "ivanantonov.com")
    add_value(doc, "https://www.linkedin.com/in/ivanantonov/")
    add_value(doc, "https://github.com/ivanpayments")
    add_field_label(doc, "Copyright")
    add_value(doc, "\u00a9 2026 Ivan Antonov. All rights reserved")


# ---------------------------------------------------------------------------
# Routing Simulator page (new)
# ---------------------------------------------------------------------------

def build_routing_simulator_page(doc):
    add_page_title(doc, "Routing Simulator landing page copy")
    add_instruction(doc, EDITOR_INSTRUCTION)

    # [RS - HERO]
    add_section_label(doc, "[RS \u2014 HERO]")
    add_field_label(doc, "Uptitle (top small line, all caps)")
    add_value(doc, "Payments infrastructure")
    add_field_label(doc, "H1 (main title)")
    add_value(
        doc,
        "A mock payment gateway for integration engineers shipping multi-PSP "
        "setups",
    )
    add_field_label(doc, "Subtitle (below title)")
    add_value(
        doc,
        "Sandboxes return the response the caller asks for. This one returns "
        "the response a merchant would actually see in production \u2014 real "
        "decline codes, 3DS flows, latency, webhook delivery",
    )
    add_field_label(doc, "Install command (mono)")
    add_value(doc, "pip install payment-routing-simulator")
    add_field_label(doc, "Stats (one per line, format: LABEL | VALUE)")
    add_value(doc, "Processor archetypes | 5")
    add_value(doc, "Synthetic transaction rows | 10,000")
    add_value(doc, "Python package | On PyPI")

    # [RS - SECTION - WHAT THIS IS]
    add_section_label(doc, "[RS \u2014 SECTION \u2014 WHAT THIS IS]")
    add_field_label(doc, "Label")
    add_value(doc, "What it is")
    add_field_label(doc, "Heading")
    add_value(doc, "One tool that emulates five kinds of payment processors")
    add_field_label(doc, "Body")
    add_value(
        doc,
        "A Python package and REST API that models five processor archetypes "
        "\u2014 global acquirer, regional bank processor, APM specialist, "
        "cross-border FX specialist, high-risk / orchestrator. Each has its "
        "own approval rates, decline codes, latency, 3DS behaviour and fees, "
        "fitted from a 10,000-row synthetic dataset",
    )

    # [RS - SECTION - THE PROBLEM IT SOLVES]
    add_section_label(doc, "[RS \u2014 SECTION \u2014 THE PROBLEM IT SOLVES]")
    add_field_label(doc, "Label")
    add_value(doc, "The problem it solves")
    add_field_label(doc, "Heading")
    add_value(
        doc,
        "Test environments don\u2019t reflect production, so routing bugs "
        "show up when customers complain",
    )
    add_field_label(
        doc,
        "Points (one per line, bold the opening phrase by ending with a "
        "period + space before the rest)",
    )
    add_value(
        doc,
        "Stripe\u2019s test mode always approves. You cannot ask \"what "
        "happens when a global acquirer soft-declines a Brazilian Visa at "
        "$300?\" \u2014 the sandbox is not wired to answer it",
    )
    add_value(
        doc,
        "Every processor\u2019s sandbox behaves differently. "
        "Regression-testing routing logic across four or five providers means "
        "maintaining four or five test harnesses",
    )
    add_value(
        doc,
        "Edge cases only surface in production. 3DS challenges, "
        "country-specific declines, soft-decline retries, webhook races "
        "\u2014 none reproduce in sandbox",
    )

    # [RS - SECTION - WHO IT IS FOR]
    add_section_label(doc, "[RS \u2014 SECTION \u2014 WHO IT IS FOR]")
    add_field_label(doc, "Label")
    add_value(doc, "Who it\u2019s for")
    add_field_label(doc, "Heading")
    add_value(
        doc,
        "Integration engineers at merchants, and QA teams shipping routing "
        "changes",
    )
    add_field_label(doc, "Body")
    add_value(
        doc,
        "Primary user: an integration engineer evaluating or maintaining a "
        "multi-PSP setup. Also useful for PSP solutions engineers, product "
        "managers comparing providers, and data engineers generating training "
        "data for retry models",
    )

    # [RS - SECTION - WHAT THE API RETURNS]
    add_section_label(doc, "[RS \u2014 SECTION \u2014 WHAT THE API RETURNS]")
    add_field_label(doc, "Label")
    add_value(doc, "What the API returns")
    add_field_label(doc, "Heading")
    add_value(doc, "Three things the sandboxes don\u2019t")
    add_field_label(doc, "Lead (italic callout)")
    add_value(
        doc,
        "Each capability is a concrete transaction path you can hit from the "
        "command line or a test suite",
    )
    add_field_label(
        doc,
        "Points (one per line, bold the opening phrase by ending with a "
        "period + space before the rest)",
    )
    add_value(
        doc,
        "Realistic decline codes by country and brand. A Brazilian Visa at "
        "$300 on the global-acquirer archetype returns ISO 8583 code 05 at "
        "the frequency that archetype actually returns it \u2014 not a 200 OK",
    )
    add_value(
        doc,
        "Full payment state machine in Postgres. PENDING to AUTHORIZED to "
        "CAPTURED or VOIDED to REFUNDED, strictly validated. Capturing a "
        "declined transaction returns 409 with the valid next states",
    )
    add_value(
        doc,
        "HMAC-signed webhooks that survive restarts. Register URL + secret; "
        "events fire with X-Signature-256 and retry on 1-2-4-8-16s backoff. "
        "Duplicate requests hit a Redis idempotency cache",
    )

    # [RS - SECTION - RUN THE REPLAY]
    add_section_label(doc, "[RS \u2014 SECTION \u2014 RUN THE REPLAY]")
    add_field_label(doc, "Label")
    add_value(doc, "Run the replay")
    add_field_label(doc, "Heading")
    add_value(
        doc,
        "Run the included replay on a committed dataset in 30 seconds",
    )
    add_field_label(doc, "Body")
    add_value(
        doc,
        "One script replays a $50M synthetic DTC apparel book through two "
        "routing configs \u2014 a single global-acquirer baseline, and a "
        "per-corridor archetype selection with soft-decline cascade, network "
        "tokens, and EEA SCA exemptions. The CSV is committed, so the output "
        "is identical on every run. Full decomposition, pattern IDs, and SQL "
        "in case_study.pdf",
    )
    add_field_label(doc, "Code block (verbatim)")
    add_value(
        doc,
        "git clone github.com/ivanpayments/payment-routing-simulator && "
        "python scripts/run_case_study.py",
    )
    add_field_label(doc, "Pull quote")
    add_value(
        doc,
        "+2.3 pp authorization, 22% soft-decline salvage within 24 hours, "
        "-18 bps effective fee. Every number traces to a CSV column",
    )
    add_field_label(doc, "Stats (one per line, format: LABEL | VALUE)")
    add_value(doc, "Book GMV replayed | $50M")
    add_value(doc, "Replay wall-clock time | 30 s")
    add_value(doc, "Data source | 10k-row synthetic CSV")

    # [RS - SECTION - TRY IT YOURSELF]
    add_section_label(doc, "[RS \u2014 SECTION \u2014 TRY IT YOURSELF]")
    add_field_label(doc, "Label")
    add_value(doc, "Try it yourself")
    add_field_label(doc, "Heading")
    add_value(doc, "Install it, hit the endpoint")
    add_field_label(doc, "Body")
    add_value(
        doc,
        "Three lines gets you a running simulator plus one comparison call:",
    )
    add_field_label(doc, "Code block (verbatim)")
    add_value(
        doc,
        "pip install payment-routing-simulator\n"
        "prs serve &\n"
        "curl -X POST localhost:8080/api/compare -d "
        "'{\"country\":\"BR\",\"brand\":\"visa\",\"amount\":300}'",
    )

    # [RS - SECTION - STACK]
    add_section_label(doc, "[RS \u2014 SECTION \u2014 STACK]")
    add_field_label(doc, "Label")
    add_value(doc, "Stack")
    add_field_label(doc, "Heading")
    add_value(
        doc,
        "FastAPI in front of Postgres, Redis, Kafka and Celery, behind Caddy",
    )
    add_field_label(doc, "Body paragraph 1")
    add_value(
        doc,
        "FastAPI validates the API key, checks Redis for idempotency, and "
        "passes a sliding-window rate limit. The engine loads archetype YAML, "
        "samples approval against country and brand modifiers, and draws "
        "latency from a log-normal fit. Every state transition writes to "
        "Postgres and publishes to Kafka. Celery workers HMAC-sign the "
        "payload and post with exponential backoff",
    )
    add_field_label(doc, "Body paragraph 2")
    add_value(
        doc,
        "The chatbot at ivanantonov.com/chatbot calls this simulator\u2019s "
        "/recommend-route endpoint to rank archetypes by expected approval, "
        "latency and fee",
    )
    add_field_label(doc, "Links (one per line, format: LABEL | URL)")
    add_value(doc, "GitHub repo | https://github.com/ivanpayments/payment-routing-simulator")
    add_value(doc, "PyPI package | https://pypi.org/project/payment-routing-simulator/")
    add_value(doc, "OpenAPI docs | /routing-simulator/api/docs")

    # [RS - CALL TO ACTION]
    add_section_label(doc, "[RS \u2014 CALL TO ACTION]")
    add_field_label(doc, "Button text")
    add_value(doc, "Read the OpenAPI docs")
    add_field_label(doc, "Primary URL")
    add_value(doc, "/routing-simulator/api/docs")
    add_field_label(doc, "Secondary text")
    add_value(doc, "pip install payment-routing-simulator")
    add_field_label(doc, "Secondary URL")
    add_value(doc, "https://pypi.org/project/payment-routing-simulator/")

    # [RS - COMPARE PANEL]
    add_section_label(doc, "[RS \u2014 COMPARE PANEL]")
    add_field_label(doc, "Status pill")
    add_value(doc, "\u25cf Live")
    add_field_label(doc, "Title")
    add_value(doc, "Compare archetypes")
    add_field_label(doc, "Subtitle")
    add_value(
        doc,
        "Pick a comparison to see the request and the ranked response the "
        "live /api/compare endpoint returns",
    )
    add_field_label(doc, "Chips label")
    add_value(doc, "Try a comparison")
    add_field_label(
        doc,
        "Chips (one per line, format: ARCHETYPE_PAIR | REQUEST)",
    )
    add_value(
        doc,
        "Global-acquirer vs regional-bank-processor | Mexican MasterCard "
        "debit at $500",
    )
    add_value(doc, "APM-specialist | Brazilian PIX payment at R$300")
    add_value(
        doc,
        "Cross-border-FX-specialist vs global-acquirer | USD\u2192EUR Visa "
        "at $120",
    )
    add_value(
        doc,
        "High-risk-orchestrator | US Visa $2,500 with 3DS v2.2 challenge",
    )
    add_value(
        doc,
        "Soft-decline cascade: global-acquirer \u2192 regional-bank-processor "
        "| Code 05 on BR Visa $300",
    )
    add_value(
        doc,
        "Webhook delivery race | Idempotency key hit on duplicate /simulate "
        "within 3s",
    )
    add_value(
        doc,
        "Commercial-card Level-3 vs Level-2 capture | US Visa $1,200 "
        "corporate card with Level-3 line-item data",
    )

    # [RS - BYLINE]
    add_section_label(doc, "[RS \u2014 BYLINE]")
    add_field_label(doc, "Byline text")
    add_value(
        doc,
        "Built by Ivan Antonov \u2014 ex-ING Payments PM, Yuno payment-ops "
        "strategy",
    )


# ---------------------------------------------------------------------------
# Driver
# ---------------------------------------------------------------------------

def main():
    doc = Document()

    # default Normal font
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    build_chatbot_page(doc)
    add_page_break(doc)
    build_portfolio_page(doc)
    add_page_break(doc)
    build_routing_simulator_page(doc)

    doc.save(str(OUT))
    print(f"Wrote: {OUT}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")


if __name__ == "__main__":
    main()
