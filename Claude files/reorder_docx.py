# -*- coding: utf-8 -*-
"""Reorder page_text.docx: insert WHY THIS MATTERS + HOW IT IS BUILT, rearrange chatbot sections,
   rename tilda-clone.html -> index.html, add Background label before the 3 portfolio bullets."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
os.chdir(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from copy import deepcopy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

d = Document('page_text.docx')
body = d.element.body

paragraphs = list(d.paragraphs)

block_starts = [(i, p.text.strip()) for i, p in enumerate(paragraphs) if p.style.name == 'Heading 1']
blocks = {}
for k, (start, name) in enumerate(block_starts):
    end = block_starts[k + 1][0] if k + 1 < len(block_starts) else len(paragraphs)
    blocks[name] = [paragraphs[i] for i in range(start, end)]

lead = paragraphs[:block_starts[0][0]]
template_block = blocks['[SECTION \u2014 WHAT THIS IS]']


def clone_block(heading1, label, heading, body_text):
    new_paras_xml = []
    texts = [heading1, 'Label', label, 'Heading', heading, 'Body', body_text]
    for i, tpl in enumerate(template_block):
        new_p = deepcopy(tpl._p)
        for r in list(new_p.findall(qn('w:r'))):
            new_p.remove(r)
        tpl_runs = template_block[i]._p.findall(qn('w:r'))
        if tpl_runs:
            r = deepcopy(tpl_runs[0])
            for t in list(r.findall(qn('w:t'))):
                r.remove(t)
            t_el = OxmlElement('w:t')
            t_el.text = texts[i]
            t_el.set(qn('xml:space'), 'preserve')
            r.append(t_el)
            new_p.append(r)
        new_paras_xml.append(new_p)
    return new_paras_xml


why_matters = clone_block(
    '[SECTION \u2014 WHY THIS MATTERS]',
    'Why this matters',
    'Today this takes days of manual work',
    ('Merchants typically dedicate three to five analysts to downloading PSP data by hand every day, '
     'track 200+ routing rules in Excel, and detect provider degradation only after hundreds of failed '
     'transactions. This chatbot turns a multi-day investigation into a two-minute conversation'),
)

how_built = clone_block(
    '[SECTION \u2014 HOW IT IS BUILT]',
    'How it is built',
    'Claude Opus writing scripts in a sandboxed environment',
    ('Claude Opus 4.6 writes and runs Python scripts in a sandboxed environment against the transaction '
     'dataset, and streams the answer back in plain English. Common questions carry pre-generated answers '
     'for instant replies. The same agent also answers on WhatsApp and SMS'),
)

# Portfolio fixes: rename filename + add Background label.
for p in paragraphs:
    if 'tilda-clone.html' in p.text:
        for run in p.runs:
            if 'tilda-clone.html' in run.text:
                run.text = run.text.replace('tilda-clone.html', 'index.html')
        if 'tilda-clone.html' in p.text:
            new_text = p.text.replace('tilda-clone.html', 'index.html')
            for r in list(p._p.findall(qn('w:r'))):
                p._p.remove(r)
            p.add_run(new_text)

for p in paragraphs:
    if 'Digital transformation at ING' in p.text and 'Background' not in p.text:
        new_text = p.text
        if '\n- Digital transformation at ING' in new_text:
            new_text = new_text.replace(
                '\n- Digital transformation at ING',
                '\nBackground\n- Digital transformation at ING', 1)
        elif '- Digital transformation at ING' in new_text:
            new_text = new_text.replace(
                '- Digital transformation at ING',
                'Background\n- Digital transformation at ING', 1)
        if new_text != p.text:
            for r in list(p._p.findall(qn('w:r'))):
                p._p.remove(r)
            p.add_run(new_text)
        break

# Split CHAT UI LABELS block into pure-chat-UI vs portfolio chunk.
chat_ui_paras = blocks['[CHAT UI LABELS]']
portfolio_start_idx = None
for i, p in enumerate(chat_ui_paras):
    t = p.text
    if '=====' in t or '[PORTFOLIO' in t or 'Portfolio landing page copy' in t:
        portfolio_start_idx = i
        break
if portfolio_start_idx is None:
    chatui_chunk = chat_ui_paras
    portfolio_chunk = []
else:
    chatui_chunk = chat_ui_paras[:portfolio_start_idx]
    portfolio_chunk = chat_ui_paras[portfolio_start_idx:]

new_chatbot_order = [
    '[HERO]',
    '[SECTION \u2014 WHAT THIS IS]',
    '__WHYMATTERS__',
    '[SECTION \u2014 WHO IT IS FOR]',
    '[SECTION \u2014 WHY IT EXISTS]',
    '[SECTION \u2014 TRANSACTIONS]',
    '__HOWBUILT__',
    '[CALL TO ACTION]',
    '[CHAT PANEL]',
    '[SUGGESTED QUESTIONS]',
    '[CHAT UI LABELS]',
]


def paras_of(name):
    if name == '[CHAT UI LABELS]':
        return chatui_chunk
    return blocks[name]


ordered_chatbot = []
for name in new_chatbot_order:
    if name == '__WHYMATTERS__':
        ordered_chatbot.extend(why_matters)
    elif name == '__HOWBUILT__':
        ordered_chatbot.extend(how_built)
    else:
        ordered_chatbot.extend([p._p for p in paras_of(name)])

final_order = [p._p for p in lead] + ordered_chatbot + [p._p for p in portfolio_chunk]
seen = set(id(e) for e in final_order)

all_p = list(body.findall(qn('w:p')))
for p_xml in all_p:
    if id(p_xml) not in seen:
        final_order.append(p_xml)
        seen.add(id(p_xml))

sectPr = body.find(qn('w:sectPr'))

# Remove every existing <w:p> from body
for p_xml in all_p:
    body.remove(p_xml)

# Reinsert in final order before sectPr
if sectPr is not None:
    idx = list(body).index(sectPr)
    for p_xml in final_order:
        body.insert(idx, p_xml)
        idx += 1
else:
    for p_xml in final_order:
        body.append(p_xml)

d.save('page_text.docx')

# Verify
d2 = Document('page_text.docx')
for i, p in enumerate(d2.paragraphs):
    if p.style.name == 'Heading 1':
        print(f'{i:3d}: {p.text}')

# Confirm portfolio fixes
for p in d2.paragraphs:
    if 'index.html' in p.text:
        print('PORTFOLIO-FILE: OK')
        break
for p in d2.paragraphs:
    if 'Background' in p.text and 'Digital transformation' in p.text:
        print('PORTFOLIO-BACKGROUND: OK')
        break
